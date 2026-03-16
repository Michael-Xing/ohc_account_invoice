"""ES/PP验证计划书（Word）填充器"""

from copy import deepcopy
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.infrastructure.template_service import TemplateFillerStrategy

logger = logging.getLogger(__name__)


class VerificationPlanFiller(TemplateFillerStrategy):
    """ES/PP验证计划书专用填充器（docx）"""

    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path, language: Optional[str] = None) -> bool:
        """填充验证计划书模板，保留 run 样式进行占位符替换"""
        self._set_language(language)
        non_empty_fields = [k for k, v in parameters.items() if v]
        logger.info("[VerificationPlanFiller] 填充字段: %s", non_empty_fields)
        try:
            doc = Document(template_path)

            test_names = self._normalize_list_field(parameters.get("test_names"))
            if test_names:
                self._fill_list_column(doc, "{{test_name}}", test_names)

            requirements_and_standards = self._normalize_list_field(parameters.get("requirements_and_standards"))
            if requirements_and_standards:
                self._fill_list_column(doc, "{{requirements_and_standards}}", requirements_and_standards)

            flat_parameters = self._flatten_parameters(parameters)
            if not test_names:
                flat_parameters["test_name"] = self._missing_text()
            if not requirements_and_standards:
                flat_parameters["requirements_and_standards"] = self._missing_text()
            self._fallback_text_replace(doc, flat_parameters)

            doc.save(output_path)
            return True
        except Exception as e:
            logger.error("验证计划书模板填充失败: %s", str(e), exc_info=True)
            return False

    # ------------------------------------------------------------------
    # 列表字段 → 表格列填充
    # ------------------------------------------------------------------

    def _normalize_list_field(self, value: Any) -> List[str]:
        if value is None:
            return []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        if isinstance(value, str):
            text = value.strip()
            return [text] if text else []
        return []

    def _fill_list_column(self, doc: Document, placeholder: str, values: List[str]) -> None:
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if placeholder not in (cell.text or ""):
                        continue
                    self._fill_column_from_anchor(table, row_idx, col_idx, values)
                    return
        logger.warning("[VerificationPlanFiller] 未找到 %s 占位符，跳过列填充", placeholder)

    def _fill_column_from_anchor(self, table, start_row: int, col_idx: int, values: List[str]) -> None:
        text_style = self._extract_cell_style(table.cell(start_row, col_idx))

        available = len(table.rows) - start_row
        if len(values) > available:
            for _ in range(len(values) - available):
                self._clone_last_row(table)

        for offset, value in enumerate(values):
            row_idx = start_row + offset
            self._set_cell_value(table.cell(row_idx, col_idx), value, text_style)

    # ------------------------------------------------------------------
    # 兜底占位符替换（正文 + 表格 + 页眉页脚）
    # ------------------------------------------------------------------

    def _fallback_text_replace(self, doc: Document, flat_parameters: Dict[str, str]) -> None:
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, flat_parameters)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, flat_parameters)

        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_in_paragraph(paragraph, flat_parameters)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_in_paragraph(paragraph, flat_parameters)

    def _replace_in_paragraph(self, paragraph, flat_parameters: Dict[str, str]) -> None:
        """在段落中替换占位符，拆分 run 使得只有填充值着色，前后缀保持原样。

        无论占位符在单个 run 内还是跨 run，都按"前缀（原色）+ 值（蓝色）+ 后缀（原色）"拆分。
        """
        full_text = paragraph.text or ""
        if not full_text:
            return

        ideal_text = full_text
        for key, value in flat_parameters.items():
            ideal_text = ideal_text.replace(f"{{{{{key}}}}}", value)

        if ideal_text == full_text:
            return

        placeholder_re = re.compile(
            r"\{\{(" + "|".join(re.escape(k) for k in flat_parameters) + r")\}\}"
        )

        if not paragraph.runs:
            paragraph.text = ideal_text
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(115, 159, 215)
            return

        font = self._extract_run_font_from_run(paragraph.runs[0])

        parts = placeholder_re.split(full_text)
        segments: List[Tuple[str, bool]] = []
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    segments.append((part, False))
            else:
                value = flat_parameters.get(part, "")
                if value:
                    segments.append((value, True))

        for run in list(paragraph.runs):
            run.clear()

        for text, is_filled in segments:
            new_run = paragraph.add_run(text)
            self._apply_font(new_run, font)
            if is_filled:
                new_run.font.color.rgb = RGBColor(115, 159, 215)

    # ------------------------------------------------------------------
    # 单元格写值
    # ------------------------------------------------------------------

    def _set_cell_value(self, cell, text: str, style: Dict = None) -> None:
        """写入单元格文本：清空所有 run 后用 add_run 写入，并恢复字体 + 段落对齐"""
        paragraph = cell.paragraphs[0] if cell.paragraphs else None
        if paragraph and paragraph.runs:
            if not style:
                style = self._extract_cell_style(cell)
            for run in list(paragraph.runs):
                run.clear()
            new_run = paragraph.add_run(text)
            self._apply_font(new_run, style)
            new_run.font.color.rgb = RGBColor(115, 159, 215)
            if style.get("alignment") is not None:
                paragraph.alignment = style["alignment"]
        else:
            cell.text = text
            p = cell.paragraphs[0] if cell.paragraphs else None
            if p:
                if style and style.get("alignment") is not None:
                    p.alignment = style["alignment"]
                if p.runs:
                    self._apply_font(p.runs[0], style)
                    p.runs[0].font.color.rgb = RGBColor(115, 159, 215)

    # ------------------------------------------------------------------
    # 行操作
    # ------------------------------------------------------------------

    def _clone_last_row(self, table) -> None:
        new_tr = deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
        for cell in table.rows[-1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""

    # ------------------------------------------------------------------
    # 字体工具
    # ------------------------------------------------------------------

    def _extract_cell_style(self, cell) -> Dict:
        """从单元格首段提取字体属性 + 段落对齐方式"""
        if not cell.paragraphs:
            return {}
        paragraph = cell.paragraphs[0]
        style: Dict[str, Any] = {}
        style["alignment"] = paragraph.alignment
        if paragraph.runs:
            style.update(self._extract_run_font_from_run(paragraph.runs[0]))
        return style

    def _extract_run_font_from_run(self, run) -> Dict:
        props: Dict[str, Any] = {}
        if run.font.name:
            props["name"] = run.font.name
        if run.font.size:
            props["size"] = run.font.size
        props["bold"] = run.font.bold
        props["italic"] = run.font.italic
        if run.font.color and run.font.color.rgb is not None:
            props["color_rgb"] = run.font.color.rgb
        elif run.font.color and run.font.color.theme_color is not None:
            props["color_theme"] = run.font.color.theme_color
        return props

    def _apply_font(self, run, font: Dict = None) -> None:
        """统一设置 run 字体（字形相关），不主动修改颜色"""
        if not font:
            return
        if font.get("name"):
            run.font.name = font["name"]
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font["name"])
        if font.get("size"):
            run.font.size = font["size"]
        if font.get("bold") is not None:
            run.font.bold = bool(font["bold"])
        if font.get("italic") is not None:
            run.font.italic = bool(font["italic"])

    # ------------------------------------------------------------------
    # 参数工具
    # ------------------------------------------------------------------

    def _flatten_parameters(self, parameters: Dict[str, Any]) -> Dict[str, str]:
        flat: Dict[str, str] = {}
        list_fields = {"test_names", "requirements_and_standards"}
        for key, value in parameters.items():
            if key in list_fields:
                continue
            if isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    if isinstance(sub_val, (str, int, float)):
                        if isinstance(sub_val, str) and sub_val.strip() == "":
                            flat[sub_key] = self._missing_text()
                        else:
                            flat[sub_key] = str(sub_val)
                    elif sub_val is None:
                        flat[sub_key] = self._missing_text()
            elif isinstance(value, (str, int, float)):
                if isinstance(value, str) and value.strip() == "":
                    flat[key] = self._missing_text()
                else:
                    flat[key] = str(value)
            elif value is None:
                flat[key] = self._missing_text()
        return flat
