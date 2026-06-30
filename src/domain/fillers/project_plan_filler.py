"""项目计划书（Word）填充器"""

from copy import deepcopy
import io
import logging
import re
import urllib.request
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.infrastructure.template_service import TemplateFillerStrategy

logger = logging.getLogger(__name__)

# Markdown 图片语法：alt 固定为 image，括号内可为 http(s) URL 或相对路径
_MD_IMAGE_RE = re.compile(r"!\[image\]\(([^)\s]+)\)", re.IGNORECASE)
# 裸 URL（不与 Markdown 图片括号区域重叠时才视为图片）
_URL_RE = re.compile(r"https?://[^\s\)\]\'\"]+")
# HTML 换行标签，后续统一转为 \n 再按段落/软换行渲染
_BR_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
_IMAGE_DOWNLOAD_FAILED_PREFIX = {
    "zh": "图片下载失败，需人工确认。地址：",
    "ja": "画像のダウンロードに失敗しました。要手動確認。URL：",
    "en": "Image download failed; manual confirmation required. URL: ",
}
_HTML_TABLE_RE = re.compile(r"<table[^>]*>.*?</table>", re.DOTALL | re.IGNORECASE)
_FILL_COLOR = RGBColor(115, 159, 215)


class ProjectPlanFiller(TemplateFillerStrategy):
    """项目计划书专用填充器（docx）

    填充流程：遍历各字段 → 判断是否富内容 → 分支渲染 → 最后批量纯文本替换
    """

    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path, language: Optional[str] = None) -> bool:
        """填充项目计划书模板"""
        self._set_language(language)
        non_empty_fields = [k for k, v in parameters.items() if v]
        logger.info("[ProjectPlanFiller] 填充字段: %s", non_empty_fields)
        try:
            doc = Document(template_path)
            self._fill_all_fields(doc, parameters)
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error("项目计划书模板填充失败: %s", str(e), exc_info=True)
            return False

    def _fill_all_fields(self, doc: Document, parameters: Dict[str, Any]) -> None:
        """遍历字段：富内容走混排渲染，纯文本收集后统一替换"""
        flat_parameters: Dict[str, str] = {}

        for key, raw in self._build_param_lookup(parameters).items():
            if self._is_rich_content(raw):
                # 富内容分支：结构性替换（段落/表格/图片）
                parts = self._build_content_parts(raw)
                if parts and self._replace_rich_placeholder(doc, f"{{{{{key}}}}}", parts):
                    logger.info(
                        "[ProjectPlanFiller] 字段 %s 富内容已处理，插入 %d 张图片",
                        key,
                        self._count_images_in_parts(parts),
                    )
                    continue
                # 未能结构性替换（无片段或未找到占位符）→ 降级为纯文本 fallback
                flat_parameters[key] = self._param_value_to_text(raw)
            else:
                # 纯文本分支：收集后整文档一次性替换
                flat_parameters[key] = self._param_value_to_text(raw)

        self._fallback_text_replace(doc, flat_parameters)

    # ------------------------------------------------------------------
    # 富内容解析与占位符替换
    # ------------------------------------------------------------------

    def _build_content_parts(self, raw: Any) -> List[Dict[str, Any]]:
        """将字段值转为 text / table / image 片段列表"""
        # 值为 URL 列表时：每项视为一张独立图片，按顺序排列
        if isinstance(raw, list):
            segments = self._parse_image_segments(raw)
            parts: List[Dict[str, Any]] = []
            for seg in segments:
                if seg["type"] == "image":
                    parts.append({"type": "image", "content": seg["url"]})
                else:
                    parts.append({"type": "text", "content": seg["content"]})
            return parts
        # 字符串：逐行扫描，拆出文本块、Markdown/HTML 表格、独立图片行
        text = str(raw).strip()
        return self._parse_mixed_content(text) if text else []

    def _replace_rich_placeholder(
        self, doc: Document, placeholder: str, parts: List[Dict[str, Any]]
    ) -> bool:
        """定位占位符并渲染富内容，找到并替换返回 True"""
        # 优先在正文段落中查找（占位符通常独占一段）
        for paragraph in list(doc.paragraphs):
            if placeholder not in (paragraph.text or ""):
                continue
            # 继承占位符段落的字体，用于图片下载失败时的兜底文字
            ref_font = self._extract_run_font_from_run(paragraph.runs[0]) if paragraph.runs else {}
            parent = paragraph._element.getparent()
            idx = list(parent).index(paragraph._element)
            parent.remove(paragraph._element)
            self._render_mixed_at_position(doc, parent, idx, parts, ref_font)
            return True

        # 段落未命中时，在模板已有表格的单元格内查找
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder not in (cell.text or ""):
                        continue
                    self._clear_cell(cell)
                    self._render_mixed_into_cell(doc, cell, parts)
                    return True
        return False

    def _count_images_in_parts(self, parts: List[Dict[str, Any]]) -> int:
        """统计片段中的图片数量（含 text 片段内的 inline 图片）"""
        count = 0
        for part in parts:
            if part["type"] == "image":
                count += 1
            elif part["type"] == "text":
                # 表格单元格或段落内嵌的 ![image](url) 也计入
                count += sum(
                    1 for seg in self._parse_image_segments(part["content"])
                    if seg["type"] == "image"
                )
        return count

    def _is_rich_content(self, value: Any) -> bool:
        """检测字段值是否含 Markdown / HTML / 图片，需走富内容渲染"""
        if value is None:
            return False
        # 多 URL 列表：每张图单独占一块
        if isinstance(value, list):
            return any(
                str(item).strip().startswith(("http://", "https://"))
                for item in value
            )

        text = str(value)
        if not text.strip():
            return False

        # --- HTML：含常见块级/行内标签即视为富内容 ---
        if re.search(r"<\s*(table|tr|td|th|img|br|p|div|span)\b", text, re.IGNORECASE):
            return True
        # --- 图片：Markdown 语法、裸 URL、或图文混排 ---
        if _MD_IMAGE_RE.search(text):
            return True
        if any(seg["type"] == "image" for seg in self._parse_image_segments(text)):
            return True
        # --- 表格：整段 HTML 表，或混排中的 Markdown 管道表 ---
        if self._is_markdown_table(text) or self._contains_pipe_table(text):
            return True
        # --- Markdown 结构：标题 / 列表 / 加粗（需 _render_text_to_elements 渲染）---
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                return True
            if re.match(r"^[-*+]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
                return True
            if "**" in stripped:
                return True
        return False

    def _contains_pipe_table(self, text: str) -> bool:
        """扫描文本中是否含 Markdown 管道表（表头行 + 分隔行含 | 与 ---）"""
        lines = text.splitlines()
        for i in range(len(lines) - 1):
            cur = lines[i].strip()
            nxt = lines[i + 1].strip()
            # 例：| A | B |  下一行  |---|---|
            if "|" in cur and "|" in nxt and re.search(r"[-:]+", nxt):
                return True
        return False

    def _build_param_lookup(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """展平参数字典：嵌套 dict 的子键提升到顶层，便于按 {{sub_key}} 匹配"""
        lookup: Dict[str, Any] = {}
        for key, value in parameters.items():
            if isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    lookup[sub_key] = sub_val
            else:
                lookup[key] = value
        return lookup

    def _parse_image_segments(self, value: Any) -> List[Dict[str, Any]]:
        """将字段值拆分为按顺序排列的 text / image 片段（用于段落内图文混排）"""
        if value is None:
            return []
        if isinstance(value, list):
            segments: List[Dict[str, Any]] = []
            for item in value:
                url = str(item).strip().strip("'\"")
                if url.startswith("http://") or url.startswith("https://"):
                    segments.append({"type": "image", "url": url})
            return segments

        text = str(value)
        if not text.strip():
            return []

        # 第一步：收集所有 Markdown 图片 ![image](url) 在原文中的位置
        spans: List[Tuple[int, int, str]] = []
        for match in _MD_IMAGE_RE.finditer(text):
            spans.append((match.start(), match.end(), match.group(1)))

        # 第二步：收集裸 URL，但跳过已被 Markdown 图片括号占用的区间
        occupied = [(start, end) for start, end, _ in spans]
        for match in _URL_RE.finditer(text):
            start, end = match.start(), match.end()
            # 完全落在某段 Markdown 图片语法内部 → 忽略（避免重复识别）
            if any(start >= o_start and end <= o_end for o_start, o_end in occupied):
                continue
            # 与已有 span 部分重叠 → 忽略
            if any(not (end <= o_start or start >= o_end) for o_start, o_end in occupied):
                continue
            spans.append((start, end, match.group(0)))

        # 第三步：按起始位置排序，切分出 text / image 交替片段
        spans.sort(key=lambda item: item[0])
        segments = []
        pos = 0
        for start, end, url in spans:
            if start > pos:
                segments.append({"type": "text", "content": text[pos:start]})
            segments.append({"type": "image", "url": url})
            pos = end
        if pos < len(text):
            segments.append({"type": "text", "content": text[pos:]})

        # 过滤掉空 text 片段
        return [seg for seg in segments if seg["type"] != "text" or seg["content"]]

    def _append_filled_inline(self, paragraph, text: str, ref_font: Dict[str, Any]) -> None:
        """填充值样式：继承模板字体 + 蓝色，支持 **加粗**"""
        pos = 0
        for match in re.finditer(r"\*\*(.+?)\*\*", text):
            if match.start() > pos:
                run = paragraph.add_run(text[pos:match.start()])
                self._apply_font(run, ref_font)
                run.font.color.rgb = _FILL_COLOR
            run = paragraph.add_run(match.group(1))
            self._apply_font(run, ref_font)
            run.font.color.rgb = _FILL_COLOR
            run.font.bold = True
            pos = match.end()
        if pos < len(text):
            run = paragraph.add_run(text[pos:])
            self._apply_font(run, ref_font)
            run.font.color.rgb = _FILL_COLOR

    def _create_scaled_image_paragraph(self, doc: Document, content: bytes):
        """创建按页宽等比缩放的图片段落（含左右缩进）"""
        section = doc.sections[0]
        # 可用宽度 = 页宽 - 左右页边距 - 段落左右缩进（与模板版式对齐）
        available_width = (
            section.page_width
            - section.left_margin
            - section.right_margin
            - Cm(0.74)
            - Cm(1.48)
        )
        available_width_cm = available_width / 360000.0  # EMU → cm
        if available_width_cm <= 0:
            available_width_cm = 10.0

        target_width_cm = available_width_cm
        target_height_cm = None
        try:
            from PIL import Image

            img_stream = io.BytesIO(content)
            img = Image.open(img_stream)
            img_width, img_height = img.size
            img_stream.close()
            if img_width > 0:
                # 96 DPI 假设下像素转 cm，再按页宽等比缩放高度
                img_width_cm = img_width * 2.54 / 96.0
                img_height_cm = img_height * 2.54 / 96.0
                scale_ratio = available_width_cm / img_width_cm
                target_width_cm = available_width_cm
                target_height_cm = img_height_cm * scale_ratio
        except Exception:
            # PIL 不可用时仅按宽度插入，高度由 Word 自动推算
            target_height_cm = None

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.right_indent = Cm(1.48)
        run = paragraph.add_run()
        image_stream = io.BytesIO(content)
        if target_height_cm is not None:
            run.add_picture(image_stream, width=Cm(target_width_cm), height=Cm(target_height_cm))
        else:
            run.add_picture(image_stream, width=Cm(target_width_cm))
        return paragraph

    def _download_image(self, url: str) -> bytes:
        try:
            req = urllib.request.Request(
                url,
                headers={"User-Agent": "ohc-account-invoice/1.0 (+python urllib)"},
                method="GET",
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                status = getattr(resp, "status", None)
                if status is not None and int(status) != 200:
                    return b""
                return resp.read()
        except Exception:
            return b""

    def _format_image_fallback(self, url: str) -> str:
        prefix = _IMAGE_DOWNLOAD_FAILED_PREFIX.get(self._language, _IMAGE_DOWNLOAD_FAILED_PREFIX["zh"])
        return prefix + url

    # ------------------------------------------------------------------
    # 混合 Markdown / HTML 解析
    # ------------------------------------------------------------------

    def _parse_mixed_content(self, text: str) -> List[Dict[str, Any]]:
        """将混排文本拆分为 text / table / image 片段（支持 Markdown 管道表与 HTML 表格）"""
        if not text:
            return []
        # 含 HTML 表格时先按 <table>...</table> 切块（支持多行、标签分行、与文字混排）
        if re.search(r"<\s*table\b", text, re.IGNORECASE):
            return self._parse_mixed_content_with_html_tables(text)
        return self._parse_mixed_content_lines(text)

    def _parse_mixed_content_with_html_tables(self, text: str) -> List[Dict[str, Any]]:
        """从文本中提取所有 HTML 表格块，表前/表后内容继续按行解析"""
        parts: List[Dict[str, Any]] = []
        last_end = 0
        for match in _HTML_TABLE_RE.finditer(text):
            before = text[last_end:match.start()]
            if before.strip():
                parts.extend(self._parse_mixed_content_lines(before))
            parts.append({"type": "table", "content": match.group(0).strip()})
            last_end = match.end()
        remaining = text[last_end:]
        if remaining.strip():
            parts.extend(self._parse_mixed_content_lines(remaining))
        return parts

    def _parse_mixed_content_lines(self, text: str) -> List[Dict[str, Any]]:
        """逐行解析：Markdown 管道表、独立图片行、普通文本（不含 HTML table 块）"""
        if not text:
            return []

        parts: List[Dict[str, Any]] = []
        current_text_lines: List[str] = []

        lines = text.splitlines()
        i = 0

        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()

            # --- 类型 A：独立占一行的图片（行内还有文字则走默认 text 分支，由 _parse_image_segments 拆分）---
            image_url = self._extract_image_url(line_stripped)
            if image_url and self._is_standalone_image_line(line_stripped):
                if current_text_lines:
                    text_content = "\n".join(current_text_lines)
                    if text_content.strip():
                        parts.append({"type": "text", "content": text_content})
                    current_text_lines = []
                parts.append({"type": "image", "content": image_url})
                i += 1
                continue

            # --- 类型 B：Markdown 管道表（表头 + --- 分隔行 + 数据行）---
            if i + 1 < len(lines):
                next_line_stripped = lines[i + 1].strip()
                if ("|" in line_stripped and "|" in next_line_stripped
                        and re.search(r"[-:]+", next_line_stripped)):
                    if current_text_lines:
                        text_content = "\n".join(current_text_lines)
                        if text_content.strip():
                            parts.append({"type": "text", "content": text_content})
                        current_text_lines = []

                    table_lines = [line, lines[i + 1]]
                    i += 2
                    while i < len(lines):
                        current_line = lines[i]
                        current_line_stripped = current_line.strip()
                        if not current_line_stripped:
                            if i + 1 < len(lines) and "|" in lines[i + 1].strip():
                                table_lines.append(current_line)
                                i += 1
                                continue
                            break
                        if "|" in current_line_stripped:
                            table_lines.append(current_line)
                            i += 1
                        else:
                            break

                    table_content = "\n".join(table_lines)
                    if self._is_markdown_table(table_content):
                        parts.append({"type": "table", "content": table_content})
                    else:
                        current_text_lines.extend(table_lines)
                    continue

            current_text_lines.append(line)
            i += 1

        if current_text_lines:
            text_content = "\n".join(current_text_lines)
            if text_content.strip():
                parts.append({"type": "text", "content": text_content})

        return parts

    def _is_markdown_table(self, text: str) -> bool:
        """检测是否为 Markdown 管道表或 HTML <table>"""
        if not text:
            return False
        if re.search(r"<table", text, re.IGNORECASE):
            return True

        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) < 2:
            return False

        header_line = lines[0]
        separator_line = lines[1]
        if "|" not in header_line or "|" not in separator_line:
            return False
        if header_line.count("|") < 2:  # 至少两列才有两个 |
            return False
        if not re.search(r"[-:]+", separator_line):
            return False
        # 分隔行去掉 | 和空格后应含连续的 - 或 :
        separator_clean = separator_line.replace("|", "").replace(" ", "")
        return bool(re.search(r"[-:]{2,}", separator_clean))

    def _extract_image_url(self, line: str) -> Optional[str]:
        """从单行提取图片 URL（Markdown ![...](url) 或 HTML <img src>）"""
        md_match = re.search(r"!\[.*?\]\((.*?)\)", line)
        if md_match:
            return md_match.group(1)
        html_match = re.search(r'<img[^>]+src=["\'](.*?)["\']', line, re.IGNORECASE)
        if html_match:
            return html_match.group(1)
        return None

    def _is_standalone_image_line(self, line: str) -> bool:
        """判断该行是否仅为图片（去掉图片语法后无其它文字）"""
        if not self._extract_image_url(line):
            return False
        rest = re.sub(r"!\[.*?\]\([^)]*\)", "", line)
        rest = re.sub(r"<img[^>]+>", "", rest, flags=re.IGNORECASE)
        return not rest.strip()

    def _strip_html_tags(self, text: str) -> str:
        """移除 HTML 标签并解码常见实体（<br> 应在此之前归一化为 \\n）"""
        text = re.sub(r"<[^>]+>", "", text)
        return (
            text.replace("&nbsp;", " ")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&amp;", "&")
            .replace("&quot;", '"')
        )

    def _prepare_text_content(self, text: str) -> str:
        """富文本片段预处理：先将 <br> 转为换行，再剥离其余 HTML 标签"""
        return self._strip_html_tags(self._normalize_line_breaks(text))

    # ------------------------------------------------------------------
    # 混合内容渲染
    # ------------------------------------------------------------------

    def _render_mixed_at_position(
        self,
        doc: Document,
        parent,
        insert_idx: int,
        parts: List[Dict[str, Any]],
        ref_font: Optional[Dict[str, Any]] = None,
    ) -> None:
        """在文档块级位置依次渲染 text / table / image 片段"""
        cur = insert_idx
        ref_font = ref_font or {}
        for part in parts:
            if part["type"] == "table":
                headers, rows, merge_info = self._parse_table_content(part["content"])
                if headers:
                    cur += self._insert_table_at(doc, parent, cur, headers, rows, merge_info)
            elif part["type"] == "image":
                # 独立图片行：下载后插入缩放段落，失败则写多语言兜底文字
                added, _ = self._insert_image_block_at(
                    doc, parent, cur, part["content"], ref_font, markdown_style=True
                )
                cur += added
            else:
                # text 片段：先剥离 HTML，再判断段内是否还有 inline 图片
                text_content = self._prepare_text_content(part["content"])
                segments = self._parse_image_segments(text_content)
                if any(seg["type"] == "image" for seg in segments):
                    cur += self._render_text_with_images_at_position(
                        doc, parent, cur, segments, ref_font, markdown_style=True
                    )
                else:
                    # 纯文本：按 Markdown 标题/列表/加粗渲染为多个段落元素
                    elements = self._render_text_to_elements(doc, text_content)
                    for el in elements:
                        parent.insert(cur, el)
                        cur += 1

    def _render_mixed_into_cell(self, doc: Document, cell, parts: List[Dict[str, Any]]) -> None:
        """在模板表格单元格内依次渲染 text / table / image 片段（可嵌套子表格）"""
        for part in parts:
            if part["type"] == "table":
                headers, rows, merge_info = self._parse_table_content(part["content"])
                if headers:
                    self._insert_table_into_cell(doc, cell, headers, rows, merge_info)
            elif part["type"] == "image":
                self._insert_image_into_cell(doc, cell, part["content"])
            else:
                text_content = self._prepare_text_content(part["content"])
                segments = self._parse_image_segments(text_content)
                if any(seg["type"] == "image" for seg in segments):
                    self._render_text_with_images_into_cell(doc, cell, segments)
                else:
                    for raw_line in text_content.splitlines():
                        line = raw_line.strip()
                        if not line:
                            continue
                        p = cell.add_paragraph("")
                        self._append_markdown_inline(p, line)

    def _render_text_with_images_at_position(
        self,
        doc: Document,
        parent,
        insert_idx: int,
        segments: List[Dict[str, Any]],
        ref_font: Dict[str, Any],
        markdown_style: bool = False,
    ) -> int:
        """渲染 text/image 混排片段，返回插入的元素数量"""
        cur = insert_idx
        for segment in segments:
            if segment["type"] == "text":
                elements = self._render_text_to_elements(
                    doc, self._prepare_text_content(segment["content"])
                )
                for el in elements:
                    parent.insert(cur, el)
                    cur += 1
            else:
                added, _ = self._insert_image_block_at(
                    doc, parent, cur, segment["url"], ref_font, markdown_style=markdown_style
                )
                cur += added
        return cur - insert_idx

    def _render_text_with_images_into_cell(
        self, doc: Document, cell, segments: List[Dict[str, Any]]
    ) -> None:
        for segment in segments:
            if segment["type"] == "text":
                for raw_line in self._prepare_text_content(segment["content"]).splitlines():
                    line = raw_line.strip()
                    if not line:
                        continue
                    p = cell.add_paragraph("")
                    self._append_markdown_inline(p, line)
            else:
                self._insert_image_into_cell(doc, cell, segment["url"])

    def _insert_image_block_at(
        self,
        doc: Document,
        parent,
        insert_idx: int,
        url: str,
        ref_font: Dict[str, Any],
        markdown_style: bool = False,
    ) -> Tuple[int, int]:
        """下载并插入单张图片（或失败兜底），返回 (新增块级元素数, 成功插入图片数)"""
        content = self._download_image(url)
        if content:
            paragraph = self._create_scaled_image_paragraph(doc, content)
            parent.insert(insert_idx, paragraph._element)
            return 1, 1

        logger.warning("图片下载失败: %s", url)
        fallback = self._format_image_fallback(url)
        paragraph = doc.add_paragraph("")
        # 富内容路径用表格字体样式；纯文本 fallback 路径用蓝色填充样式
        if markdown_style:
            self._append_markdown_inline(paragraph, fallback)
        else:
            self._append_filled_inline(paragraph, fallback, ref_font)
        parent.insert(insert_idx, paragraph._element)
        return 1, 0

    def _insert_image_into_cell(self, doc: Document, cell, url: str) -> None:
        """在表格单元格内插入单张图片（宽度撑满单元格可用区域）"""
        content = self._download_image(url)
        paragraph = cell.add_paragraph("")
        if content:
            run = paragraph.add_run()
            section = doc.sections[0]
            available_width = section.page_width - section.left_margin - section.right_margin
            available_width_cm = max(available_width / 360000.0, 1.0)
            image_stream = io.BytesIO(content)
            run.add_picture(image_stream, width=Cm(available_width_cm))
            return

        logger.warning("图片下载失败: %s", url)
        self._append_markdown_inline(paragraph, self._format_image_fallback(url))

    # ------------------------------------------------------------------
    # 表格解析与插入（Markdown 管道表 + HTML 表格）
    # ------------------------------------------------------------------

    def _parse_table_content(
        self, table_text: str
    ) -> Tuple[List[str], List[List[str]], List[Dict[str, Any]]]:
        """解析表格内容为 (表头, 数据行, 合并信息)；Markdown 表 merge_info 为空"""
        if re.search(r"<table", table_text, re.IGNORECASE):
            return self._parse_html_table(table_text)

        # Markdown 管道表：第 0 行表头，第 1 行分隔符，第 2 行起为数据
        lines = [line.strip() for line in table_text.splitlines() if line.strip()]
        if len(lines) < 2:
            return [], [], []
        header_line = lines[0]
        separator_line = lines[1]
        if "|" not in header_line or "|" not in separator_line:
            return [], [], []

        headers = [c.strip() for c in header_line.strip("|").split("|")]
        rows: List[List[str]] = []
        for line in lines[2:]:
            if "|" not in line:
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            # 列数不足补空，超出截断，保证与表头列数一致
            if len(cells) < len(headers):
                cells.extend([""] * (len(headers) - len(cells)))
            elif len(cells) > len(headers):
                cells = cells[: len(headers)]
            rows.append(cells)
        return headers, rows, []

    def _parse_html_table(
        self, html_text: str
    ) -> Tuple[List[str], List[List[str]], List[Dict[str, Any]]]:
        """解析 HTML 表格，支持 rowspan / colspan"""
        table_match = re.search(r"<table[^>]*>(.*?)</table>", html_text, re.DOTALL | re.IGNORECASE)
        if not table_match:
            return [], [], []

        table_content = table_match.group(1)
        tr_matches = re.findall(r"<tr[^>]*>(.*?)</tr>", table_content, re.DOTALL | re.IGNORECASE)
        if not tr_matches:
            return [], [], []

        all_rows: List[List[str]] = []
        merge_info: List[Dict[str, Any]] = []
        # rowspan_map[行索引][列索引] = (单元格文本, 剩余向下占用行数)
        rowspan_map: List[Dict[int, Tuple[str, int]]] = []

        for tr_idx, tr_content in enumerate(tr_matches):
            current_row: List[str] = []
            col_idx = 0

            # 本行开头：填入上一行 rowspan 延续下来的占位单元格
            if tr_idx < len(rowspan_map):
                for prev_col_idx in sorted(rowspan_map[tr_idx].keys()):
                    while len(current_row) < prev_col_idx:
                        current_row.append("")
                    if prev_col_idx >= len(current_row):
                        current_row.extend([""] * (prev_col_idx - len(current_row) + 1))
                    if prev_col_idx < len(current_row) and current_row[prev_col_idx] == "":
                        value, remaining = rowspan_map[tr_idx][prev_col_idx]
                        current_row[prev_col_idx] = value
                        if remaining > 1:
                            if tr_idx + 1 >= len(rowspan_map):
                                rowspan_map.append({})
                            rowspan_map[tr_idx + 1][prev_col_idx] = (value, remaining - 1)

            for cell_match in re.finditer(
                r"<t[hd]([^>]*)>(.*?)</t[hd]>", tr_content, re.DOTALL | re.IGNORECASE
            ):
                attrs = cell_match.group(1)
                inner_html = cell_match.group(2)
                rowspan_attr = re.search(r'rowspan=["\']?(\d+)', attrs, re.IGNORECASE)
                colspan_attr = re.search(r'colspan=["\']?(\d+)', attrs, re.IGNORECASE)
                rowspan = int(rowspan_attr.group(1)) if rowspan_attr else 1
                colspan = int(colspan_attr.group(1)) if colspan_attr else 1
                cell_text = self._strip_html_tags(inner_html).strip()

                while col_idx < len(current_row) and current_row[col_idx] != "":
                    col_idx += 1
                while len(current_row) < col_idx:
                    current_row.append("")

                cell_col = col_idx
                if rowspan > 1 or colspan > 1:
                    merge_info.append({
                        "row": tr_idx,
                        "col": cell_col,
                        "rowspan": rowspan,
                        "colspan": colspan,
                    })

                current_row.append(cell_text)
                col_idx += 1
                for _ in range(colspan - 1):
                    current_row.append("")
                    col_idx += 1

                if rowspan > 1:
                    if tr_idx + 1 >= len(rowspan_map):
                        rowspan_map.append({})
                    for offset in range(1, rowspan):
                        target_row = tr_idx + offset
                        while target_row >= len(rowspan_map):
                            rowspan_map.append({})
                        rowspan_map[target_row][cell_col] = (cell_text, rowspan - offset)

            if current_row:
                all_rows.append(current_row)

        if not all_rows:
            return [], [], []

        # 各行对齐到相同列数
        max_cols = max(len(row) for row in all_rows)
        for row in all_rows:
            if len(row) < max_cols:
                row.extend([""] * (max_cols - len(row)))

        # 约定：HTML 第一行 tr 作为表头
        return all_rows[0], all_rows[1:], merge_info

    def _fill_table_cell(self, doc: Document, cell, val: str) -> None:
        """填充表格单元格：支持行内 Markdown 与 ![image](url)"""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
        segments = self._parse_image_segments(val)
        if any(seg["type"] == "image" for seg in segments):
            # 含图片时需清空默认段落再逐段插入
            cell.text = ""
            for para in list(cell.paragraphs):
                p_el = para._element
                p_el.getparent().remove(p_el)
            self._render_text_with_images_into_cell(doc, cell, segments)
            return
        paragraph.text = ""
        self._append_markdown_inline(paragraph, self._prepare_text_content(val))

    def _insert_table_at(
        self,
        doc: Document,
        parent,
        insert_idx: int,
        headers: List[str],
        rows: List[List[str]],
        merge_info: Optional[List[Dict[str, Any]]] = None,
    ) -> int:
        """在文档中插入表格并应用合并，返回插入的元素数量"""
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        for i, h in enumerate(headers):
            run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
            self._apply_table_font(run, bold=True)
        self._apply_header_row_style(tbl.rows[0])
        for r_idx, vals in enumerate(rows, start=1):
            for c_idx, val in enumerate(vals):
                self._fill_table_cell(doc, tbl.rows[r_idx].cells[c_idx], val)
        self._apply_table_border(tbl)
        # 先填内容再合并，避免 merge 后写入异常
        if merge_info:
            self._apply_word_table_merges(tbl, merge_info)
        parent.insert(insert_idx, tbl._element)
        return 1

    def _insert_table_into_cell(
        self,
        doc: Document,
        cell,
        headers: List[str],
        rows: List[List[str]],
        merge_info: Optional[List[Dict[str, Any]]] = None,
    ) -> None:
        """在单元格内嵌套表格"""
        tbl = cell.add_table(rows=1 + len(rows), cols=len(headers))
        for i, h in enumerate(headers):
            run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
            self._apply_table_font(run, bold=True)
        self._apply_header_row_style(tbl.rows[0])
        for r_idx, vals in enumerate(rows, start=1):
            for c_idx, val in enumerate(vals):
                self._fill_table_cell(doc, tbl.rows[r_idx].cells[c_idx], val)
        self._apply_table_border(tbl)
        if merge_info:
            self._apply_word_table_merges(tbl, merge_info)

    def _apply_word_table_merges(self, table, merge_info: List[Dict[str, Any]]) -> None:
        """按 merge_info 合并 Word 表格单元格（rowspan / colspan）"""
        for merge in merge_info:
            row = merge["row"]
            col = merge["col"]
            rowspan = merge["rowspan"]
            colspan = merge["colspan"]
            if rowspan <= 1 and colspan <= 1:
                continue
            try:
                top_left = table.rows[row].cells[col]
                bottom_right = table.rows[row + rowspan - 1].cells[col + colspan - 1]
                top_left.merge(bottom_right)
            except Exception as exc:
                # 合并失败不阻断整体填充，仅记录警告
                logger.warning(
                    "Word 表格合并失败 row=%s col=%s: %s", row, col, exc
                )

    # ------------------------------------------------------------------
    # 文本渲染
    # ------------------------------------------------------------------

    def _render_text_to_elements(self, doc: Document, text: str) -> List:
        """将文本渲染为段落元素（标题 / 列表 / 加粗）；text 应已通过 _prepare_text_content 预处理"""
        elements = []
        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if not line:
                p = doc.add_paragraph("")
                elements.append(p._element)
                continue

            # Markdown 标题 # / ## / ...
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                content = line[level:].strip()
                p = doc.add_paragraph("")
                style_name = f"Heading {level}"
                if style_name in doc.styles:
                    p.style = doc.styles[style_name]
                self._append_markdown_inline(p, content)
                elements.append(p._element)
                continue

            # 无序列表 - / *
            if re.match(r"^[-*]\s+", line):
                content = re.sub(r"^[-*]\s+", "", line)
                p = doc.add_paragraph(style="List Bullet" if "List Bullet" in doc.styles else None)
                self._append_markdown_inline(p, content)
                elements.append(p._element)
                continue

            # 有序列表 1. / 2. ...
            if re.match(r"^\d+\.\s+", line):
                content = re.sub(r"^\d+\.\s+", "", line)
                p = doc.add_paragraph(style="List Number" if "List Number" in doc.styles else None)
                self._append_markdown_inline(p, content)
                elements.append(p._element)
                continue

            p = doc.add_paragraph("")
            self._append_markdown_inline(p, line)
            elements.append(p._element)

        return elements

    def _normalize_line_breaks(self, text: str) -> str:
        """将 <br> / <br/> 归一化为 \\n"""
        return _BR_RE.sub("\n", text)

    def _append_markdown_inline(self, paragraph, text: str) -> None:
        """行内 Markdown：\\n / <br> 软换行，**text** 加粗（使用表格字体色）"""
        text = self._normalize_line_breaks(text)
        lines = text.split("\n")
        for idx, line in enumerate(lines):
            if idx > 0:
                paragraph.add_run().add_break()  # Word 段内换行
            if line:
                self._append_markdown_bold_runs(paragraph, line)

    def _append_markdown_bold_runs(self, paragraph, text: str) -> None:
        """单行内 **加粗** 渲染（table 字体）"""
        pos = 0
        for match in re.finditer(r"\*\*(.+?)\*\*", text):
            if match.start() > pos:
                run = paragraph.add_run(text[pos:match.start()])
                self._apply_table_font(run)
            run = paragraph.add_run(match.group(1))
            self._apply_table_font(run, bold=True)
            pos = match.end()
        if pos < len(text):
            run = paragraph.add_run(text[pos:])
            self._apply_table_font(run)

    # ------------------------------------------------------------------
    # 单元格/样式工具
    # ------------------------------------------------------------------

    def _clear_cell(self, cell) -> None:
        """清空单元格全部段落，保留一个空段落供后续写入"""
        cell.text = ""
        for para in list(cell.paragraphs):
            p_el = para._element
            p_el.getparent().remove(p_el)
        cell.add_paragraph("")

    def _apply_table_font(self, run, bold: bool = False) -> None:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(115, 159, 215)
        run.font.bold = bold

    def _apply_header_row_style(self, row) -> None:
        for cell in row.cells:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9D9D9")
            cell._tc.get_or_add_tcPr().append(shading)

    def _apply_table_border(self, table) -> None:
        tbl_pr = table._element.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._element.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            borders.append(b)
        tbl_pr.append(borders)

    # ------------------------------------------------------------------
    # 兜底占位符替换（正文 + 表格 + 页眉页脚）
    # ------------------------------------------------------------------

    def _fallback_text_replace(self, doc: Document, flat_parameters: Dict[str, str]) -> None:
        """对未走富内容路径的字段，将 {{key}} 替换为纯文本（蓝色）"""
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, flat_parameters)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, flat_parameters)

        # 页眉页脚同样支持占位符，但不处理图片/表格混排
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_in_paragraph(paragraph, flat_parameters)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_in_paragraph(paragraph, flat_parameters)

    def _replace_in_paragraph(self, paragraph, flat_parameters: Dict[str, str]) -> None:
        """在段落中替换占位符，拆分 run 使得只有填充值着色，前后缀保持原样。

        无论占位符在单个 run 内还是跨 run，都按"前缀（原色）+ 值（蓝色）+ 后缀（原色）"拆分。
        """
        full_text = paragraph.text or ""
        if not full_text:
            return

        ideal_text = full_text
        for key, value in flat_parameters.items():
            ideal_text = ideal_text.replace(f"{{{{{key}}}}}", value)

        if ideal_text == full_text:
            return

        placeholder_re = re.compile(
            r"\{\{(" + "|".join(re.escape(k) for k in flat_parameters) + r")\}\}"
        )

        # 对整段拼合文本做拆分，确保跨 run 场景也能精确定位
        parts = placeholder_re.split(full_text)
        # parts: [普通文字, key, 普通文字, key, ...]
        segments: List[Tuple[str, bool]] = []
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    segments.append((part, False))
            else:
                value = flat_parameters.get(part, "")
                if value:
                    segments.append((value, True))

        font = self._extract_run_font_from_run(paragraph.runs[0]) if paragraph.runs else {}

        for run in list(paragraph.runs):
            run.clear()

        for text, is_filled in segments:
            new_run = paragraph.add_run(text)
            self._apply_font(new_run, font)
            if is_filled:
                new_run.font.color.rgb = RGBColor(115, 159, 215)

    # ------------------------------------------------------------------
    # 单元格写值
    # ------------------------------------------------------------------

    def _set_cell_value(self, cell, text: str, style: Dict = None) -> None:
        """写入单元格文本：清空所有 run 后用 add_run 写入，并恢复字体 + 段落对齐"""
        paragraph = cell.paragraphs[0] if cell.paragraphs else None
        if paragraph and paragraph.runs:
            if not style:
                style = self._extract_cell_style(cell)
            for run in list(paragraph.runs):
                run.clear()
            new_run = paragraph.add_run(text)
            self._apply_font(new_run, style)
            new_run.font.color.rgb = RGBColor(115, 159, 215)
            if style.get("alignment") is not None:
                paragraph.alignment = style["alignment"]
        else:
            cell.text = text
            p = cell.paragraphs[0] if cell.paragraphs else None
            if p:
                if style and style.get("alignment") is not None:
                    p.alignment = style["alignment"]
                if p.runs:
                    self._apply_font(p.runs[0], style)
                    p.runs[0].font.color.rgb = RGBColor(115, 159, 215)

    # ------------------------------------------------------------------
    # 行操作
    # ------------------------------------------------------------------

    def _clone_last_row(self, table) -> None:
        new_tr = deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
        for cell in table.rows[-1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""

    # ------------------------------------------------------------------
    # 字体工具
    # ------------------------------------------------------------------

    def _extract_cell_style(self, cell) -> Dict:
        """从单元格首段提取字体属性 + 段落对齐方式"""
        if not cell.paragraphs:
            return {}
        paragraph = cell.paragraphs[0]
        style: Dict[str, Any] = {}
        style["alignment"] = paragraph.alignment
        if paragraph.runs:
            style.update(self._extract_run_font_from_run(paragraph.runs[0]))
        return style

    def _extract_run_font_from_run(self, run) -> Dict:
        props: Dict[str, Any] = {}
        if run.font.name:
            props["name"] = run.font.name
        if run.font.size:
            props["size"] = run.font.size
        props["bold"] = run.font.bold
        props["italic"] = run.font.italic
        if run.font.color and run.font.color.rgb is not None:
            props["color_rgb"] = run.font.color.rgb
        elif run.font.color and run.font.color.theme_color is not None:
            props["color_theme"] = run.font.color.theme_color
        return props

    def _apply_font(self, run, font: Dict = None) -> None:
        """统一设置 run 字体，并尽量恢复原有显式颜色"""
        if not font:
            return
        if font.get("name"):
            run.font.name = font["name"]
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font["name"])
        if font.get("size"):
            run.font.size = font["size"]
        if font.get("bold") is not None:
            run.font.bold = bool(font["bold"])
        if font.get("italic") is not None:
            run.font.italic = bool(font["italic"])
        if font.get("color_rgb") is not None:
            run.font.color.rgb = font["color_rgb"]
        elif font.get("color_theme") is not None:
            run.font.color.theme_color = font["color_theme"]

    # ------------------------------------------------------------------
    # 参数工具
    # ------------------------------------------------------------------

    def _param_value_to_text(self, value: Any) -> str:
        """将单个参数字段转为纯文本填充值（空值用缺失占位文案）"""
        if value is None:
            return self._missing_text()
        if isinstance(value, (str, int, float)):
            if isinstance(value, str) and value.strip() == "":
                return self._missing_text()
            return str(value)
        return str(value)
