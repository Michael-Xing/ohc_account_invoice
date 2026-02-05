from pathlib import Path
from typing import Any, Dict, List, Tuple

import io
import re
import tempfile

import requests
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches

from src.infrastructure.template_service import TemplateFillerStrategy


class BasicSpecificationFiller(TemplateFillerStrategy):
    """基本规格书专用填充器"""

    # 需要转换为 Word 表格的 markdown 字段
    MARKDOWN_TABLE_FIELDS = {
        "definition_term_table",
        "component_table",
        "function_table",
        "function_block_table",
        "performance_table",
    }

    # 图片列表字段
    IMAGE_LIST_FIELDS = {
        "appearance_image",
        "function_block_image",
    }

    # 派生的商品型号表字段占位符
    PRODUCT_MODEL_TABLE_FIELD = "product_model_table"
    
    # 需要行首缩进4字符的字段
    INDENT_4_CHARS_FIELDS = {
        "dimensions_and_weight",
        "power_supply",
        "use_temperature_humidity_range",
        "storage_and_transport_conditions",
        "durability",
    }

    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path) -> bool:
        """填充基本规格书模板"""
        try:
            # 预处理参数：构造商品型号表数据
            self._build_product_model_table(parameters)

            # 加载 Word 模板
            doc = Document(template_path)

            # 先在表格单元格中处理占位符（表格更适合放 markdown 表和图片）
            self._process_tables(doc, parameters)

            # 再处理普通段落中的占位符（用于 markdown 文本转段落/列表等）
            self._process_paragraphs(doc, parameters)

            # 最后做一次兜底的纯文本占位符替换，避免遗漏简单文本
            flat_parameters = self._flatten_parameters(parameters)
            self._fallback_text_replace(doc, flat_parameters)

            # 保存结果
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"基本规格书模板填充失败: {str(e)}")
            import traceback

            traceback.print_exc()
            return False

    def _build_product_model_table(self, parameters: Dict[str, Any]) -> None:
        """根据商品型号和贩卖名称构造商品型号表数据，用于 product_model_table 占位符"""
        product_model = str(parameters.get("product_model") or "").strip()
        sales_name = str(parameters.get("sales_name") or "").strip()

        if not product_model and not sales_name:
            return

        product_models = [item.strip() for item in product_model.split("/") if item.strip()]
        sales_names = [item.strip() for item in sales_name.split("/") if item.strip()]

        if not product_models and not sales_names:
            return

        rows: List[Dict[str, str]] = []
        max_len = max(len(product_models), len(sales_names))
        for idx in range(max_len):
            rows.append(
                {
                    "sales_name": sales_names[idx] if idx < len(sales_names) else "",
                    "catalog_number": "",
                    "ohq_product_model": product_models[idx] if idx < len(product_models) else "",
                    "basic_udi_di_code": "",
                    "device_category": "",
                }
            )

        parameters[self.PRODUCT_MODEL_TABLE_FIELD] = rows

    def _process_tables(self, doc: Document, parameters: Dict[str, Any]) -> None:
        """处理所有表格中的占位符：markdown表格、商品型号表、图片等"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text or ""
                    placeholders = self._extract_placeholders(text)
                    if not placeholders:
                        continue

                    for placeholder in placeholders:
                        key = placeholder
                        if key in self.MARKDOWN_TABLE_FIELDS:
                            markdown_text = str(self._get_param(parameters, key) or "").strip()
                            if markdown_text:
                                self._clear_cell(cell)
                                self._insert_markdown_table_into_cell(cell, markdown_text, merge_same_column=False)
                        elif key == "function_block_table":
                            markdown_text = str(self._get_param(parameters, key) or "").strip()
                            if markdown_text:
                                self._clear_cell(cell)
                                self._insert_markdown_table_into_cell(cell, markdown_text, merge_same_column=False)
                        elif key == "performance_table":
                            markdown_text = str(self._get_param(parameters, key) or "").strip()
                            if markdown_text:
                                self._clear_cell(cell)
                                self._insert_markdown_table_into_cell(cell, markdown_text, merge_same_column=False)
                        elif key == self.PRODUCT_MODEL_TABLE_FIELD:
                            rows = parameters.get(self.PRODUCT_MODEL_TABLE_FIELD) or []
                            if rows:
                                self._clear_cell(cell)
                                self._insert_product_model_table(cell, rows)
                        elif key in self.IMAGE_LIST_FIELDS:
                            images = self._normalize_image_urls(self._get_param(parameters, key))
                            if images:
                                self._clear_cell(cell)
                                self._insert_images_into_cell(cell, images)
                                # 占位符已通过 clear_cell 清除，无需再次清理
                                continue

                        # 占位符处理完，避免重复文本留在单元格中（只在段落级别清理，保留格式）
                        if f"{{{{{key}}}}}" in cell.text:
                            for paragraph in cell.paragraphs:
                                if paragraph.text and f"{{{{{key}}}}}" in paragraph.text:
                                    # 在 run 级别替换，保留格式
                                    if paragraph.runs:
                                        first_run = paragraph.runs[0]
                                        font_name = first_run.font.name
                                        font_size = first_run.font.size
                                        is_bold = first_run.font.bold
                                        is_italic = first_run.font.italic
                                        font_color = first_run.font.color
                                        
                                        new_text = paragraph.text.replace(f"{{{{{key}}}}}", "")
                                        
                                        # 清空所有 runs
                                        for run in list(paragraph.runs):
                                            run.clear()
                                        
                                        # 用原格式创建新文本
                                        if new_text:
                                            new_run = paragraph.add_run(new_text)
                                            if font_name:
                                                new_run.font.name = font_name
                                            if font_size:
                                                new_run.font.size = font_size
                                            new_run.font.bold = is_bold
                                            new_run.font.italic = is_italic
                                            if font_color:
                                                new_run.font.color = font_color
                                    else:
                                        # 没有 run，直接替换
                                        paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", "")

    def _process_paragraphs(self, doc: Document, parameters: Dict[str, Any]) -> None:
        """处理文档中普通段落的占位符，将 markdown 文本转换为段落/列表/加粗等结构"""
        for paragraph in list(doc.paragraphs):
            text = paragraph.text or ""
            placeholders = self._extract_placeholders(text)
            if not placeholders:
                continue

            # 处理每个占位符
            for placeholder in placeholders:
                key = placeholder
                parent_element = paragraph._element
                parent = parent_element.getparent()

                # 检查是否是独占一行的占位符（用于表格和复杂内容）
                is_standalone = len(placeholders) == 1 and text.strip() == f"{{{{{key}}}}}"

                # 1）如果是 markdown 表格类字段或派生的商品型号表，占位符独占一行时直接在该位置插入 Word 表格
                if (key in self.MARKDOWN_TABLE_FIELDS or key == self.PRODUCT_MODEL_TABLE_FIELD) and is_standalone:
                    # 在删除之前保存插入位置
                    insert_idx = parent.index(parent_element)
                    # 删除占位符段落
                    parent.remove(parent_element)

                    if key == self.PRODUCT_MODEL_TABLE_FIELD:
                        rows = parameters.get(self.PRODUCT_MODEL_TABLE_FIELD) or []
                        if rows:
                            self._insert_product_model_table_at_block(doc, parent, insert_idx, rows)
                    else:
                        value = self._get_param(parameters, key)
                        markdown_text = str(value or "").strip()
                    if markdown_text:
                        # 不再进行列合并
                        merge_same = False
                        self._insert_markdown_table_at_block(doc, parent, insert_idx, markdown_text, merge_same)
                    break  # 处理完这个段落，跳出循环
                elif key in self.IMAGE_LIST_FIELDS:
                    # 3）图片列表字段：在段落位置插入图片（无论是否独占一行）
                    image_param = self._get_param(parameters, key)
                    images = self._normalize_image_urls(image_param)
                    print(f"调试: 字段 {key} 的原始值: {image_param}, 提取的图片URLs: {images}")
                    if images:
                        # 在删除之前保存插入位置
                        insert_idx = parent.index(parent_element)
                        # 删除占位符段落
                        parent.remove(parent_element)
                        # 在指定位置插入图片
                        self._insert_images_at_block(doc, parent, insert_idx, images)
                    else:
                        print(f"警告: 字段 {key} 没有找到有效的图片URL，原始值: {image_param}")
                        # 即使没有图片，也要删除占位符，避免被 _fallback_text_replace 处理
                        if is_standalone:
                            parent.remove(parent_element)
                    break  # 处理完这个段落，跳出循环
                elif is_standalone:
                    # 2）普通 markdown 文本：按段落/标题/列表处理（仅当独占一行时）
                    value = self._get_param(parameters, key)
                    markdown_text = str(value or "").strip()

                    # 在删除之前保存插入位置
                    insert_idx = parent.index(parent_element)
                    # 删除原始占位符段落
                    parent.remove(parent_element)

                    if markdown_text:
                        # 判断是否需要行首缩进4字符
                        indent_4_chars = key in self.INDENT_4_CHARS_FIELDS
                        self._render_markdown_block(doc, parent, insert_idx, markdown_text, indent_4_chars=indent_4_chars)
                    break  # 处理完这个段落，跳出循环

    def _fallback_text_replace(self, doc: Document, flat_parameters: Dict[str, str]) -> None:
        """兜底的纯文本占位符替换，避免遗漏简单字符串场景，同时保留原有格式"""

        def has_placeholder(text: str) -> bool:
            """检查文本是否包含占位符"""
            if not text:
                return False
            # 排除图片列表字段和表格字段的占位符
            excluded_fields = self.IMAGE_LIST_FIELDS | self.MARKDOWN_TABLE_FIELDS | {self.PRODUCT_MODEL_TABLE_FIELD}
            for key in flat_parameters.keys():
                if key not in excluded_fields and f"{{{{{key}}}}}" in text:
                    return True
            return False

        def replace_in_runs(paragraph, flat_parameters: Dict[str, str]) -> None:
            """在 run 级别替换占位符，保留原有格式"""
            if not paragraph.text or not has_placeholder(paragraph.text):
                return

            # 合并所有 runs 的文本
            full_text = paragraph.text
            new_text = full_text
            for key, value in flat_parameters.items():
                new_text = new_text.replace(f"{{{{{key}}}}}", value)

            if full_text == new_text:
                return

            # 如果有 runs，保留第一个 run 的格式
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # 获取格式属性
                font_name = first_run.font.name
                font_size = first_run.font.size
                is_bold = first_run.font.bold
                is_italic = first_run.font.italic
                font_color = first_run.font.color

                # 清空所有 runs
                for run in list(paragraph.runs):
                    run.clear()

                # 用第一个 run 的格式创建新文本，并应用新字体样式
                if new_text:
                    new_run = paragraph.add_run(new_text)
                    # 应用新的字体样式（微软雅黑 10号 #7F7F7F）
                    self._apply_font(new_run, bold=is_bold)
                    # 应用首行缩进2字符
                    paragraph.paragraph_format.first_line_indent = Cm(0.74)
            else:
                # 没有 run，直接设置（会使用默认格式）
                paragraph.text = new_text
                # 应用首行缩进2字符
                paragraph.paragraph_format.first_line_indent = Cm(0.74)
                # 应用字体样式
                for run in paragraph.runs:
                    self._apply_font(run, bold=False)

        # 处理段落：只替换包含占位符的段落
        for paragraph in doc.paragraphs:
            replace_in_runs(paragraph, flat_parameters)

        # 处理表格单元格：只替换包含占位符的单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_runs(paragraph, flat_parameters)

    def _insert_product_model_table(self, cell, rows: List[Dict[str, str]]) -> None:
        """在单元格中插入商品型号表"""
        headers = ["销售名称", "Catalog number", "OHQ商品型式名", "Basic UDI-DI code", "医疗器械类别分类"]
        table = cell.add_table(rows=1 + len(rows), cols=len(headers))

        # 设置表头
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            self._apply_font(run, bold=True)
        self._apply_header_row_style(table.rows[0])

        # 填充数据行
        for row_idx, data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = data.get("sales_name", "")
            row_cells[1].text = data.get("catalog_number", "")
            row_cells[2].text = data.get("ohq_product_model", "")
            row_cells[3].text = data.get("basic_udi_di_code", "")
            row_cells[4].text = data.get("device_category", "")

        # 应用正文字体样式
        for row in table.rows:
            for cell_item in row.cells:
                for paragraph in cell_item.paragraphs:
                    for run in paragraph.runs:
                        self._apply_font(run, bold=run.bold)

    def _insert_product_model_table_at_block(self, doc: Document, parent, insert_idx: int, rows: List[Dict[str, str]]) -> None:
        """在文档块级位置插入商品型号表"""
        headers = ["销售名称", "Catalog number", "OHQ商品型式名", "Basic UDI-DI code", "医疗器械类别分类"]
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))

        # 设置表头
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            self._apply_font(run, bold=True)
        self._apply_header_row_style(table.rows[0])

        # 填充数据行
        for row_idx, data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = data.get("sales_name", "")
            row_cells[1].text = data.get("catalog_number", "")
            row_cells[2].text = data.get("ohq_product_model", "")
            row_cells[3].text = data.get("basic_udi_di_code", "")
            row_cells[4].text = data.get("device_category", "")

        # 应用正文字体样式
        for row in table.rows:
            for cell_item in row.cells:
                for paragraph in cell_item.paragraphs:
                    for run in paragraph.runs:
                        self._apply_font(run, bold=run.bold)
        
        # 应用表格样式（边框、换行处理）
        self._apply_table_style(table)
        
        # 在表格前后添加缩进段落（首尾各缩进2字符）
        # 先插入表格前的缩进段落
        indent_para_before = doc.add_paragraph()
        indent_para_before.paragraph_format.left_indent = Cm(0.74)
        indent_para_before.paragraph_format.right_indent = Cm(0.74)
        
        # 插入表格
        tbl_element = table._element
        parent.insert(insert_idx, indent_para_before._element)
        parent.insert(insert_idx + 1, tbl_element)
        
        # 插入表格后的缩进段落（尾部缩进4字符）
        indent_para_after = doc.add_paragraph()
        indent_para_after.paragraph_format.left_indent = Cm(0.74)
        indent_para_after.paragraph_format.right_indent = Cm(1.48)  # 4字符 = 1.48cm
        parent.insert(insert_idx + 2, indent_para_after._element)

    def _insert_markdown_table_into_cell(self, cell, markdown_text: str, merge_same_column: bool) -> None:
        """将 markdown 表格解析后插入到单元格中"""
        headers, rows = self._parse_markdown_table(markdown_text)
        if not headers:
            return

        table = cell.add_table(rows=1 + len(rows), cols=len(headers))

        # 表头
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            self._apply_font(run, bold=True)
        self._apply_header_row_style(table.rows[0])

        # 数据行
        for row_idx, row_vals in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, val in enumerate(row_vals):
                row_cells[col_idx].text = val

        # 合并相同列内容（功能模块、性能说明等需要）
        if merge_same_column:
            self._merge_same_content_columns(table)

        # 设置所有单元格字体
        for row in table.rows:
            for cell_item in row.cells:
                for paragraph in cell_item.paragraphs:
                    for run in paragraph.runs:
                        self._apply_font(run, bold=run.bold)
        
        # 应用表格样式（边框、缩进、换行处理）
        self._apply_table_style(table)

    def _insert_markdown_table_at_block(
        self, doc: Document, parent, insert_idx: int, markdown_text: str, merge_same_column: bool
    ) -> None:
        """在文档块级位置插入 markdown 表格"""
        headers, rows = self._parse_markdown_table(markdown_text)
        if not headers:
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))

        # 表头
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            self._apply_font(run, bold=True)
        self._apply_header_row_style(table.rows[0])

        # 数据行
        for row_idx, row_vals in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, val in enumerate(row_vals):
                row_cells[col_idx].text = val

        # 合并相同列内容（功能模块、性能说明等需要）
        if merge_same_column:
            self._merge_same_content_columns(table)

        # 设置所有单元格字体
        for row in table.rows:
            for cell_item in row.cells:
                for paragraph in cell_item.paragraphs:
                    for run in paragraph.runs:
                        self._apply_font(run, bold=run.bold)
        
        # 应用表格样式（边框、换行处理）
        self._apply_table_style(table)
        
        # 在表格前后添加缩进段落（首尾各缩进2字符）
        # 先插入表格前的缩进段落
        indent_para_before = doc.add_paragraph()
        indent_para_before.paragraph_format.left_indent = Cm(0.74)
        indent_para_before.paragraph_format.right_indent = Cm(0.74)
        
        # 插入表格
        tbl_element = table._element
        parent.insert(insert_idx, indent_para_before._element)
        parent.insert(insert_idx + 1, tbl_element)
        
        # 插入表格后的缩进段落（尾部缩进4字符）
        indent_para_after = doc.add_paragraph()
        indent_para_after.paragraph_format.left_indent = Cm(0.74)
        indent_para_after.paragraph_format.right_indent = Cm(1.48)  # 4字符 = 1.48cm
        parent.insert(insert_idx + 2, indent_para_after._element)

    def _insert_images_into_cell(self, cell, image_urls: List[str]) -> None:
        """将多张图片插入到同一个单元格中，按给定顺序，保持原始比例"""
        for url in image_urls:
            try:
                content = self._download_image(url)
                if not content:
                    continue
                image_stream = io.BytesIO(content)
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                # 添加图片，保持原始比例（不指定宽度和高度）
                run.add_picture(image_stream)
            except Exception:
                # 单张图片失败不中断整体处理
                continue

    def _insert_images_at_block(self, doc: Document, parent, insert_idx: int, image_urls: List[str]) -> None:
        """在文档块级位置插入多张图片，按给定顺序，宽度适配word文档，首行缩进2字符，尾部缩进4字符，保持宽高比"""
        elements_to_insert = []
        
        # 获取文档页面宽度（减去左右边距）
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        # 可用宽度 = 页面宽度 - 左边距 - 右边距 - 首行缩进(2字符) - 尾部缩进(4字符)
        available_width = page_width - left_margin - right_margin - Cm(0.74) - Cm(1.48)
        
        for url in image_urls:
            try:
                content = self._download_image(url)
                if not content:
                    print(f"警告: 图片下载失败或内容为空: {url}")
                    continue
                
                # 计算可用宽度（EMU 单位转换为 Cm）
                # 1 cm = 360000 EMU
                available_width_cm = available_width / 360000.0
                
                if available_width_cm <= 0:
                    print(f"警告: 可用宽度计算错误: {available_width_cm} cm，使用默认宽度 10cm")
                    available_width_cm = 10.0
                
                # 尝试使用 PIL 读取图片尺寸以计算缩放比例
                target_width_cm = None
                target_height_cm = None
                try:
                    from PIL import Image
                    img_stream_for_size = io.BytesIO(content)
                    img = Image.open(img_stream_for_size)
                    img_width, img_height = img.size
                    img_stream_for_size.close()  # 关闭用于读取尺寸的流
                    
                    if img_width > 0:
                        # 将图片像素尺寸转换为厘米（假设图片是 96 DPI，1英寸 = 2.54cm，1英寸 = 96像素）
                        # 1像素 = 2.54/96 cm ≈ 0.026458333 cm
                        img_width_cm = img_width * 2.54 / 96.0
                        img_height_cm = img_height * 2.54 / 96.0
                        
                        # 计算缩放比例（以可用宽度为准）
                        scale_ratio = available_width_cm / img_width_cm
                        target_width_cm = available_width_cm
                        target_height_cm = img_height_cm * scale_ratio
                    else:
                        print(f"警告: 图片宽度为0: {url}，使用默认尺寸")
                        target_width_cm = available_width_cm
                        target_height_cm = available_width_cm  # 默认正方形
                except ImportError:
                    # PIL 不可用，使用默认宽度，让 python-docx 自动计算高度
                    print(f"警告: PIL 模块不可用，使用默认宽度 {available_width_cm:.2f}cm，高度自动计算")
                    target_width_cm = available_width_cm
                    target_height_cm = None  # 不指定高度，让 python-docx 保持宽高比
                except Exception as e:
                    # 读取图片尺寸失败，使用默认尺寸
                    print(f"警告: 读取图片尺寸失败: {str(e)}，使用默认尺寸")
                    target_width_cm = available_width_cm
                    target_height_cm = None  # 不指定高度，让 python-docx 保持宽高比
                
                # 创建新的流对象用于插入图片
                image_stream_for_insert = io.BytesIO(content)
                
                paragraph = doc.add_paragraph()
                # 设置首行缩进2字符，尾部缩进4字符
                paragraph.paragraph_format.left_indent = Cm(0.74)
                paragraph.paragraph_format.right_indent = Cm(1.48)
                run = paragraph.add_run()
                # 添加图片，指定宽度和高度（保持宽高比）
                if target_height_cm is not None:
                    run.add_picture(image_stream_for_insert, width=Cm(target_width_cm), height=Cm(target_height_cm))
                else:
                    # 只指定宽度，让 python-docx 自动计算高度以保持宽高比
                    run.add_picture(image_stream_for_insert, width=Cm(target_width_cm))
                image_stream_for_insert.close()  # 关闭流
                
                elements_to_insert.append(paragraph._element)
                if target_height_cm is not None:
                    print(f"成功插入图片: {url}, 尺寸: {target_width_cm:.2f}cm x {target_height_cm:.2f}cm")
                else:
                    print(f"成功插入图片: {url}, 尺寸: {target_width_cm:.2f}cm x 自动")
            except Exception as e:
                # 单张图片失败不中断整体处理
                import traceback
                print(f"错误: 插入图片失败 {url}: {str(e)}")
                traceback.print_exc()
                continue
        
        # 将所有图片段落按顺序插入到指定位置
        for idx, element in enumerate(elements_to_insert):
            parent.insert(insert_idx + idx, element)

    def _render_markdown_block(self, doc: Document, parent, insert_idx: int, markdown_text: str, indent_4_chars: bool = False) -> None:
        """将 markdown 文本渲染为一组段落/列表，插入到指定位置"""
        lines = markdown_text.splitlines()
        
        # 记录当前插入位置，从初始位置开始
        current_idx = insert_idx
        elements_to_insert = []

        for raw_line in lines:
            line = raw_line.rstrip()
            if not line:
                # 空行插入一个普通空段落
                p = doc.add_paragraph("")
                self._apply_paragraph_style(p, indent_4_chars=indent_4_chars)
                elements_to_insert.append(p._element)
                continue

            # 标题（仅处理 # 和 ##）
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line[level:].strip()
                p = doc.add_paragraph("")
                if level == 1:
                    p.style = doc.styles["Heading 1"] if "Heading 1" in doc.styles else None
                elif level == 2:
                    p.style = doc.styles["Heading 2"] if "Heading 2" in doc.styles else None
                self._append_markdown_inline(p, text)
                elements_to_insert.append(p._element)
                continue

            # 无序列表
            if re.match(r"^[-*]\s+", line):
                text = re.sub(r"^[-*]\s+", "", line)
                p = doc.add_paragraph(style="List Bullet" if "List Bullet" in doc.styles else None)
                self._append_markdown_inline(p, text)
                elements_to_insert.append(p._element)
                continue

            # 有序列表
            if re.match(r"^\d+\.\s+", line):
                text = re.sub(r"^\d+\.\s+", "", line)
                p = doc.add_paragraph(style="List Number" if "List Number" in doc.styles else None)
                self._append_markdown_inline(p, text)
                elements_to_insert.append(p._element)
                continue

            # 普通段落
            p = doc.add_paragraph("")
            self._append_markdown_inline(p, line)
            self._apply_paragraph_style(p, indent_4_chars=indent_4_chars)
            elements_to_insert.append(p._element)

        # 将所有元素按顺序插入到指定位置
        for idx, element in enumerate(elements_to_insert):
            parent.insert(insert_idx + idx, element)

    def _append_markdown_inline(self, paragraph, text: str) -> None:
        """处理行内 markdown（目前支持 **加粗**）并写入到段落中"""
        pos = 0
        pattern = re.compile(r"\*\*(.+?)\*\*")
        for match in pattern.finditer(text):
            if match.start() > pos:
                normal_text = text[pos : match.start()]
                if normal_text:
                    run = paragraph.add_run(normal_text)
                    self._apply_font(run, bold=False)
            bold_text = match.group(1)
            run = paragraph.add_run(bold_text)
            self._apply_font(run, bold=True)
            pos = match.end()

        if pos < len(text):
            tail = text[pos:]
            if tail:
                run = paragraph.add_run(tail)
                self._apply_font(run, bold=False)

    def _apply_font(self, run, bold: bool = False) -> None:
        """统一设置字体为微软雅黑 10 号，颜色 #7F7F7F"""
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)  # #7F7F7F
        run.font.bold = bool(bold)

    def _apply_paragraph_style(self, paragraph, indent_4_chars: bool = False) -> None:
        """段落基础样式：左对齐、微软雅黑 10 号、首行缩进2字符或4字符"""
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 首行缩进：4字符字段缩进4字符（1.48cm），其他缩进2字符（0.74cm）
        if indent_4_chars:
            paragraph.paragraph_format.first_line_indent = Cm(1.48)  # 4字符 = 1.48cm
        else:
            paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 2字符 = 0.74cm
        for run in paragraph.runs:
            self._apply_font(run, bold=run.bold)

    def _apply_header_row_style(self, row) -> None:
        """表头样式：浅灰背景 + 加粗微软雅黑"""
        fill_color = "D9D9D9"
        for cell in row.cells:
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), fill_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._apply_font(run, bold=True)

    def _apply_table_style(self, table) -> None:
        """设置表格样式：首尾缩进2字符、黑色边框、单元格内文本处理换行"""
        # 设置表格首尾缩进（左右各缩进2字符 = 0.74cm）
        # 在 Word 中，表格缩进通过设置表格的左右边距实现
        tbl_pr = table._element.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._element.insert(0, tbl_pr)
        
        # 设置左缩进（2字符 = 0.74cm = 420 twips）
        # 1cm = 567 twips, 0.74cm ≈ 420 twips
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), "420")  # 420 twips = 0.74cm
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)
        
        # 设置表格边框为黑色
        tbl_borders = OxmlElement("w:tblBorders")
        border_style = "single"
        border_size = "4"  # 0.5pt
        border_color = "000000"  # 黑色
        
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), border_style)
            border.set(qn("w:sz"), border_size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), border_color)
            tbl_borders.append(border)
        
        tbl_pr.append(tbl_borders)
        
        # 处理单元格中的 <br> 换行符，转换为 Word 换行
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text or ""
                    if text and (re.search(r'<br\s*/?>', text, re.IGNORECASE)):
                        # 处理换行符
                        # 替换所有类型的 <br> 标签为换行符
                        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
                        # 清空段落
                        paragraph.clear()
                        # 按换行符分割并创建多个 run
                        lines = text.split('\n')
                        for idx, line in enumerate(lines):
                            if line:
                                run = paragraph.add_run(line)
                                self._apply_font(run, bold=False)
                            # 如果不是最后一行，添加换行符
                            if idx < len(lines) - 1:
                                paragraph.add_run().add_break()

    def _merge_same_content_columns(self, table) -> None:
        """对表中每一列内容相同的连续单元格进行纵向合并"""
        if len(table.rows) <= 2:
            return

        num_rows = len(table.rows)
        num_cols = len(table.columns)

        for col_idx in range(num_cols):
            start_row = 1
            while start_row < num_rows:
                current_text = table.cell(start_row, col_idx).text
                end_row = start_row + 1
                while end_row < num_rows and table.cell(end_row, col_idx).text == current_text:
                    end_row += 1
                if end_row - start_row > 1 and current_text:
                    top_cell = table.cell(start_row, col_idx)
                    bottom_cell = table.cell(end_row - 1, col_idx)
                    top_cell.merge(bottom_cell)
                start_row = end_row

    def _parse_markdown_table(self, markdown_text: str) -> Tuple[List[str], List[List[str]]]:
        """解析 markdown 表格文本为表头和数据行"""
        lines = [line.strip() for line in markdown_text.splitlines() if line.strip()]
        if len(lines) < 2:
            return [], []

        header_line = lines[0]
        separator_line = lines[1]
        if "|" not in header_line or "|" not in separator_line:
            return [], []

        headers = [cell.strip() for cell in header_line.strip("|").split("|")]
        rows: List[List[str]] = []
        for line in lines[2:]:
            if "|" not in line:
                continue
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            # 对齐列数
            if len(cells) < len(headers):
                cells.extend([""] * (len(headers) - len(cells)))
            elif len(cells) > len(headers):
                cells = cells[: len(headers)]
            rows.append(cells)

        return headers, rows

    def _normalize_image_urls(self, value: Any) -> List[str]:
        """将入参中的图片字段整理为 URL 列表，并做简单清洗"""
        urls: List[str] = []
        if isinstance(value, list):
            candidates = value
        elif isinstance(value, str):
            candidates = [value]
        else:
            candidates = []

        pattern = re.compile(r"https?://[^\s)]+")
        for item in candidates:
            if not item:
                continue
            text = str(item)
            # 从文本中抽取 URL
            found = pattern.findall(text)
            if found:
                urls.extend(found)
            else:
                # 文本本身就是 URL
                urls.append(text.strip())
        return urls

    def _download_image(self, url: str) -> bytes:
        """下载图片内容，失败时返回空字节串"""
        try:
            resp = requests.get(url, timeout=10)
            if resp.status_code == 200:
                return resp.content
        except Exception:
            return b""
        return b""

    def _extract_placeholders(self, text: str) -> List[str]:
        """从文本中提取 {{var}} 形式的占位符变量名"""
        if not text:
            return []
        return re.findall(r"\{\{([^}]+)\}\}", text)

    def _get_param(self, parameters: Dict[str, Any], key: str) -> Any:
        """支持从嵌套对象中取值的简单访问（仅支持一层嵌套）"""
        if key in parameters:
            return parameters[key]

        # 针对已知嵌套结构的简单展开映射
        nested_mappings = {
            "power_supply": ("service_environment_conditions", "power_supply"),
            "use_temperature_humidity_range": ("service_environment_conditions", "use_temperature_humidity_range"),
            "storage_and_transport_conditions": ("service_environment_conditions", "storage_and_transport_conditions"),
            "durability": ("service_environment_conditions", "durability"),
            "definitions_of_basic_safety": ("safety_protection_info", "definitions_of_basic_safety"),
            "device_classification": ("safety_protection_info", "device_classification"),
            "equipment_safety_protection_and_warnings": (
                "safety_protection_info",
                "equipment_safety_protection_and_warnings",
            ),
            "safety_protection": ("safety_protection_info", "safety_protection"),
            "safety_warning": ("safety_protection_info", "safety_warning"),
            "biological_alarms": ("safety_protection_info", "biological_alarms"),
            "technical_alarms": ("safety_protection_info", "technical_alarms"),
            "default_equipment_setting": ("various_settings", "default_equipment_setting"),
            "date_time_settings": ("various_settings", "date_time_settings"),
            "maintenance": ("maintenance_and_disposal", "maintenance"),
            "disposal": ("maintenance_and_disposal", "disposal"),
        }

        if key in nested_mappings:
            parent_key, child_key = nested_mappings[key]
            parent_val = parameters.get(parent_key)
            if isinstance(parent_val, dict):
                return parent_val.get(child_key)
            # Pydantic 模型在传入时会被转换为 dict
        return None

    def _flatten_parameters(self, parameters: Dict[str, Any]) -> Dict[str, str]:
        """将嵌套参数展平为简单字典，用于兜底文本替换"""
        flat: Dict[str, str] = {}
        # 排除图片列表字段和表格字段，避免被兜底替换处理
        excluded_fields = self.IMAGE_LIST_FIELDS | self.MARKDOWN_TABLE_FIELDS | {self.PRODUCT_MODEL_TABLE_FIELD}
        
        for key, value in parameters.items():
            # 跳过图片列表字段和表格字段
            if key in excluded_fields:
                continue
            if isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    if isinstance(sub_val, (str, int, float)):
                        flat[sub_key] = str(sub_val)
            elif isinstance(value, (str, int, float)):
                flat[key] = str(value)
        return flat

    def _clear_cell(self, cell) -> None:
        """清空单元格内容"""
        cell.text = ""
        # 清除所有段落，保留一个空段落作为插入锚点
        for para in list(cell.paragraphs):
            p_element = para._element
            p_element.getparent().remove(p_element)
        cell.add_paragraph("")

    def _insert_after(self, parent, paragraph, previous_element):
        """在指定元素之后插入段落，如果 previous_element 为空则附加到末尾"""
        new_element = paragraph._element
        if previous_element is None:
            parent.append(new_element)
        else:
            parent.insert(parent.index(previous_element) + 1, new_element)
        return new_element

