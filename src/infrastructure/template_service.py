"""模板填充服务模块（已迁移到 infrastructure 层）"""

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from openpyxl import load_workbook

from src.config import settings


class TemplateFillerStrategy(ABC):
    """模板填充策略抽象基类"""

    # 空值兜底文本（按 language 选择）
    _MISSING_TEXT_BY_LANGUAGE = {
        "zh": "AI未检索到，需人工确认",
        "ja": "AIで検索できませんでした。要手動確認",
        "en": "AI could not retrieve this; manual confirmation required",
    }

    def __init__(self) -> None:
        self._language: str = "zh"

    def _set_language(self, language: Optional[str]) -> None:
        """设置当前语言"""
        lang = (language or "").strip().lower()
        if lang in self._MISSING_TEXT_BY_LANGUAGE:
            self._language = lang
        else:
            self._language = "zh"

    def _missing_text(self) -> str:
        """获取当前语言对应的空值兜底文本"""
        return self._MISSING_TEXT_BY_LANGUAGE.get(self._language, self._MISSING_TEXT_BY_LANGUAGE["zh"])

    def _is_missing_text(self, value: Any) -> bool:
        """判断值是否为空值（None、空字符串或任何语言的兜底文本）"""
        if value is None:
            return True
        if not isinstance(value, str):
            return False
        s = value.strip()
        if s == "":
            return True
        return s in set(self._MISSING_TEXT_BY_LANGUAGE.values())

    @abstractmethod
    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path, language: Optional[str] = None) -> bool:
        """填充模板"""
        pass
    
    def _replace_placeholders(self, text: str, parameters: Dict[str, Any]) -> str:
        """替换文本中的占位符"""
        if not text or not isinstance(text, str):
            return text
        
        def replace_double_brace(match):
            placeholder = match.group(1).strip()
            if placeholder in parameters:
                value = parameters[placeholder]
                # 如果值为空，使用语言对应的兜底文本
                if self._is_missing_text(value):
                    return self._missing_text()
                return str(value)
            return match.group(0)
        
        text = re.sub(r'\{\{([^}]+)\}\}', replace_double_brace, text)
        
        def replace_single_brace(match):
            placeholder = match.group(1).strip()
            if placeholder in parameters:
                value = parameters[placeholder]
                # 如果值为空，使用语言对应的兜底文本
                if self._is_missing_text(value):
                    return self._missing_text()
                return str(value)
            return match.group(0)
        
        text = re.sub(r'\{([^}]+)\}', replace_single_brace, text)
        
        return text


class ExcelTemplateFiller(TemplateFillerStrategy):
    """Excel模板填充策略"""

    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path, language: Optional[str] = None) -> bool:
        try:
            # 设置语言
            self._set_language(language)

            workbook = load_workbook(template_path)
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                for row in worksheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            cell.value = self._replace_placeholders(cell.value, parameters)
            workbook.save(output_path)
            return True
        except Exception as e:
            print(f"Excel模板填充失败: {str(e)}")
            return False


class WordTemplateFiller(TemplateFillerStrategy):
    """Word模板填充策略"""

    def fill_template(self, template_path: Path, parameters: Dict[str, Any], output_path: Path, language: Optional[str] = None) -> bool:
        try:
            # 设置语言
            self._set_language(language)

            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    paragraph.text = self._replace_placeholders(paragraph.text, parameters)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            cell.text = self._replace_placeholders(cell.text, parameters)
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text:
                            paragraph.text = self._replace_placeholders(paragraph.text, parameters)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text:
                            paragraph.text = self._replace_placeholders(paragraph.text, parameters)
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Word模板填充失败: {str(e)}")
            return False