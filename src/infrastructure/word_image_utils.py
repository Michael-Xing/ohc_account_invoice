"""Word 文档图片工具：URL 解析、下载与插入。

供各 Word filler 复用，与具体账票模板解耦。

占位符替换主流程（replace_placeholder_with_images）：
  1. 定位 {{key}}（正文 → 表格 → 页眉/页脚）
  2. 找到占位符后，逐 URL 下载（不在定位前批量下载）
  3. 至少一张成功 → 删占位符 → 插入图片 → 部分失败则追加兜底文本
  4. 全部失败 → 保留原段落/单元格，占位符原地替换为兜底文本
"""

import ast
import io
import logging
import re
import urllib.request
from typing import Any, Callable, List, Optional, Tuple

from docx import Document
from docx.shared import Cm

ReplaceInParagraphFn = Callable[[Any, str, str], None]
ReplaceInCellFn = Callable[[Any, str, str], None]
InsertBlockParagraphFn = Callable[[Document, Any, int, str, Any], None]
InsertCellParagraphFn = Callable[[Any, str, Any], None]

logger = logging.getLogger(__name__)

_URL_PATTERN = re.compile(r"https?://[^\s\)\]\'\"]+")

_IMAGE_DOWNLOAD_FAILED_PREFIX = {
    "zh": "图片下载失败，需人工确认。地址：",
    "ja": "画像のダウンロードに失敗しました。要手動確認。URL：",
    "en": "Image download failed; manual confirmation required. URL: ",
}


def format_image_download_fallback(image_urls: List[str], language: str = "zh") -> str:
    """生成图片下载失败时的兜底文本（含错误提示与 URL）"""
    prefix = _IMAGE_DOWNLOAD_FAILED_PREFIX.get(language, _IMAGE_DOWNLOAD_FAILED_PREFIX["zh"])
    return prefix + "\n".join(image_urls)

def normalize_image_urls(value: Any) -> List[str]:
    """将入参中的图片字段整理为 URL 列表，并做简单清洗
    
    支持多种输入格式：
    - 列表: ['url1', 'url2']
    - 字符串格式的列表: "['url1', 'url2']"
    - 单个 URL 字符串: "https://..."
    - 包含多个 URL 的字符串: "url1 url2" 或 "url1, url2"
    """
    urls: List[str] = []
    
    if value is None:
        return urls
    
    # Step 1: 将入参统一转为候选列表
    if isinstance(value, list):
        candidates = value
    elif isinstance(value, str):
        # Step 1a: 尝试 ast 解析字符串列表
        try:
            parsed = ast.literal_eval(value)
            if isinstance(parsed, list):
                candidates = parsed
            else:
                candidates = [value]
        except (ValueError, SyntaxError):
            found = _URL_PATTERN.findall(value)
            if found:
                candidates = found
            else:
                candidates = [value]
    else:
        candidates = []

    for item in candidates:
        if not item:
            continue
        text = str(item).strip().strip("'\"")
        found = _URL_PATTERN.findall(text)
        if found:
            urls.extend(found)
        elif text.startswith("http://") or text.startswith("https://"):
            urls.append(text)
    
    return list(dict.fromkeys(urls))


def pure_image_urls(value: Any) -> List[str]:
    """value 为纯 URL 或 URL 列表时返回 URL，否则返回空列表（只解析一次）"""
    if value is None:
        return []
    if isinstance(value, str) and value.strip() == "":
        return []

    urls = normalize_image_urls(value)
    if not urls:
        return []

    remainder = str(value)
    for url in urls:
        remainder = remainder.replace(url, "")
    remainder = re.sub(r"[\s,'\"\[\]]+", "", remainder)
    return urls if remainder == "" else []


def is_image_content(value: Any) -> bool:
    """判断入参是否应作为图片 URL 处理（整段内容为 URL 或 URL 列表）。

    仅当解析出的 URL 占满全部内容时返回 True；
    如「详见 https://... 文档」这类混合文本返回 False，仍走普通文本替换。
    """
    return bool(pure_image_urls(value))


def download_image(url: str) -> bytes:
    """下载单张图片，失败时返回空字节串（由调用方判断成功与否）"""
    try:
        req = urllib.request.Request(
            url,
            headers={
                "User-Agent": "ohc-account-invoice/1.0 (+python urllib)",
            },
            method="GET",
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            status = getattr(resp, "status", None)
            if status is not None and int(status) != 200:
                return b""
            return resp.read()
    except Exception:
        return b""


def _fetch_image_contents(image_urls: List[str]) -> Tuple[List[bytes], List[str]]:
    """定位占位符后逐 URL 下载，返回 (成功字节列表, 失败 URL 列表)。

    注意：仅在已找到占位符后调用，不在定位前批量预下载。
    """
    ok_contents: List[bytes] = []
    failed_urls: List[str] = []
    for url in image_urls:
        content = download_image(url)
        if content:
            ok_contents.append(content)
        else:
            failed_urls.append(url)
            logger.warning("图片下载失败: %s", url)
    return ok_contents, failed_urls


def _insert_image_contents_into_cell(cell, image_contents: List[bytes]) -> int:
    """将已下载的图片字节插入单元格（不再发起网络请求），返回成功插入的数量"""
    inserted = 0
    for content in image_contents:
        try:
            image_stream = io.BytesIO(content)
            paragraph = cell.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_stream)
            inserted += 1
        except Exception:
            continue
    return inserted


def insert_images_into_cell(cell, image_urls: List[str]) -> int:
    """将多张图片插入到同一个单元格中，返回成功插入的数量"""
    ok_contents, _ = _fetch_image_contents(image_urls)
    return _insert_image_contents_into_cell(cell, ok_contents)


def _insert_image_contents_at_block(
    doc: Document, parent, insert_idx: int, image_contents: List[bytes]
) -> int:
    """在文档块级位置插入已下载的图片（不再发起网络请求），返回成功插入的数量"""
    elements_to_insert = []

    # Step 1: 计算页面可用宽度（扣除页边距与段落缩进）
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    available_width = page_width - left_margin - right_margin - Cm(0.74) - Cm(1.48)

    for content in image_contents:
        try:
            available_width_cm = available_width / 360000.0
            if available_width_cm <= 0:
                print(f"警告: 可用宽度计算错误: {available_width_cm} cm，使用默认宽度 10cm")
                available_width_cm = 10.0

            # Step 2: 读取图片尺寸，按可用宽度等比缩放
            target_width_cm = None
            target_height_cm = None
            try:
                from PIL import Image
                img_stream_for_size = io.BytesIO(content)
                img = Image.open(img_stream_for_size)
                img_width, img_height = img.size
                img_stream_for_size.close()

                if img_width > 0:
                    img_width_cm = img_width * 2.54 / 96.0
                    img_height_cm = img_height * 2.54 / 96.0
                    scale_ratio = available_width_cm / img_width_cm
                    target_width_cm = available_width_cm
                    target_height_cm = img_height_cm * scale_ratio
                else:
                    print("警告: 图片宽度为0，使用默认尺寸")
                    target_width_cm = available_width_cm
                    target_height_cm = available_width_cm
            except ImportError:
                print(f"警告: PIL 模块不可用，使用默认宽度 {available_width_cm:.2f}cm，高度自动计算")
                target_width_cm = available_width_cm
                target_height_cm = None
            except Exception as e:
                print(f"警告: 读取图片尺寸失败: {str(e)}，使用默认尺寸")
                target_width_cm = available_width_cm
                target_height_cm = None

            # Step 3: 创建段落并插入图片
            image_stream_for_insert = io.BytesIO(content)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.right_indent = Cm(1.48)
            run = paragraph.add_run()
            if target_height_cm is not None:
                run.add_picture(
                    image_stream_for_insert, width=Cm(target_width_cm), height=Cm(target_height_cm)
                )
            else:
                run.add_picture(image_stream_for_insert, width=Cm(target_width_cm))
            image_stream_for_insert.close()

            elements_to_insert.append(paragraph._element)
        except Exception as e:
            logger.warning("插入图片失败: %s", e)
            continue

    # Step 4: 将段落元素插入到占位符原位置
    for idx, element in enumerate(elements_to_insert):
        parent.insert(insert_idx + idx, element)
    return len(elements_to_insert)


def insert_images_at_block(doc: Document, parent, insert_idx: int, image_urls: List[str]) -> int:
    """在文档块级位置插入多张图片，返回成功插入的数量"""
    ok_contents, _ = _fetch_image_contents(image_urls)
    return _insert_image_contents_at_block(doc, parent, insert_idx, ok_contents)


def clear_cell(cell) -> None:
    """清空单元格内容，保留一个空段落作为插入锚点"""
    cell.text = ""
    for para in list(cell.paragraphs):
        p_el = para._element
        p_el.getparent().remove(p_el)
    cell.add_paragraph("")


def _replace_placeholder_in_paragraph(paragraph, placeholder: str, text: str) -> None:
    """原地替换段落中的占位符文本（不删除段落本身）"""
    paragraph.text = (paragraph.text or "").replace(placeholder, text)


def _replace_placeholder_in_cell(cell, placeholder: str, text: str) -> None:
    """原地替换单元格中的占位符文本（不清空单元格）"""
    for para in cell.paragraphs:
        if placeholder in (para.text or ""):
            para.text = (para.text or "").replace(placeholder, text)
            return
    if placeholder in (cell.text or ""):
        cell.text = cell.text.replace(placeholder, text)


def _insert_fallback_paragraph(
    doc: Document,
    parent,
    insert_idx: int,
    failed_urls: List[str],
    language: str,
    *,
    insert_block_paragraph_fn: Optional[InsertBlockParagraphFn] = None,
    reference_paragraph: Any = None,
) -> None:
    """在图片段落后追加下载失败的兜底文本段落"""
    if not failed_urls:
        return
    text = format_image_download_fallback(failed_urls, language)
    if insert_block_paragraph_fn:
        insert_block_paragraph_fn(doc, parent, insert_idx, text, reference_paragraph)
        return
    paragraph = doc.add_paragraph(text)
    parent.insert(insert_idx, paragraph._element)


def _apply_images_to_paragraph(
    doc: Document,
    paragraph,
    placeholder: str,
    image_urls: List[str],
    language: str,
    *,
    replace_in_paragraph_fn: Optional[ReplaceInParagraphFn] = None,
    insert_block_paragraph_fn: Optional[InsertBlockParagraphFn] = None,
) -> None:
    """处理正文/页眉/页脚中的图片占位符段落。

    Step 1: 逐 URL 下载（此时占位符仍在文档中）
    Step 2a: 至少一张成功 → 删占位符段落 → 插入图片 → 部分失败则追加兜底
    Step 2b: 全部失败 → 不删段落，占位符原地替换为兜底文本
    """
    ok_contents, failed_urls = _fetch_image_contents(image_urls)
    fallback_text = format_image_download_fallback(image_urls, language)
    if ok_contents:
        # Step 2a-1: 下载成功，删除占位符段落（保留原段落供兜底文本继承字体）
        reference_paragraph = paragraph
        parent = paragraph._element.getparent()
        insert_idx = parent.index(paragraph._element)
        parent.remove(paragraph._element)
        # Step 2a-2: 在原位置插入图片
        inserted = _insert_image_contents_at_block(doc, parent, insert_idx, ok_contents)
        # Step 2a-3: 部分 URL 失败，在图片后追加兜底段落
        if failed_urls:
            _insert_fallback_paragraph(
                doc,
                parent,
                insert_idx + inserted,
                failed_urls,
                language,
                insert_block_paragraph_fn=insert_block_paragraph_fn,
                reference_paragraph=reference_paragraph,
            )
    else:
        # Step 2b: 全部失败，保留段落结构，仅替换占位符文字
        if replace_in_paragraph_fn:
            replace_in_paragraph_fn(paragraph, placeholder, fallback_text)
        else:
            _replace_placeholder_in_paragraph(paragraph, placeholder, fallback_text)


def _apply_images_to_cell(
    cell,
    placeholder: str,
    image_urls: List[str],
    language: str,
    *,
    replace_in_cell_fn: Optional[ReplaceInCellFn] = None,
    insert_cell_paragraph_fn: Optional[InsertCellParagraphFn] = None,
) -> None:
    """处理表格单元格中的图片占位符。

    Step 1: 逐 URL 下载（此时占位符仍在单元格中）
    Step 2a: 至少一张成功 → 清空单元格 → 插入图片 → 部分失败则追加兜底
    Step 2b: 全部失败 → 不清空单元格，占位符原地替换为兜底文本
    """
    ok_contents, failed_urls = _fetch_image_contents(image_urls)
    fallback_text = format_image_download_fallback(image_urls, language)
    if ok_contents:
        # Step 2a-1: 下载成功，清空单元格（保留原段落供兜底文本继承字体）
        reference_paragraph = cell.paragraphs[0] if cell.paragraphs else None
        clear_cell(cell)
        # Step 2a-2: 插入图片
        _insert_image_contents_into_cell(cell, ok_contents)
        # Step 2a-3: 部分 URL 失败，追加兜底段落
        if failed_urls:
            partial_fallback = format_image_download_fallback(failed_urls, language)
            if insert_cell_paragraph_fn:
                insert_cell_paragraph_fn(cell, partial_fallback, reference_paragraph)
            else:
                cell.add_paragraph(partial_fallback)
    else:
        # Step 2b: 全部失败，保留单元格结构，仅替换占位符文字
        if replace_in_cell_fn:
            replace_in_cell_fn(cell, placeholder, fallback_text)
        else:
            _replace_placeholder_in_cell(cell, placeholder, fallback_text)


def _apply_images_in_paragraphs(
    paragraphs,
    doc: Document,
    placeholder: str,
    image_urls: List[str],
    language: str,
    *,
    replace_in_paragraph_fn: Optional[ReplaceInParagraphFn] = None,
    insert_block_paragraph_fn: Optional[InsertBlockParagraphFn] = None,
) -> bool:
    for paragraph in list(paragraphs):
        if placeholder not in (paragraph.text or ""):
            continue
        _apply_images_to_paragraph(
            doc,
            paragraph,
            placeholder,
            image_urls,
            language,
            replace_in_paragraph_fn=replace_in_paragraph_fn,
            insert_block_paragraph_fn=insert_block_paragraph_fn,
        )
        return True
    return False


def replace_placeholder_with_images(
    doc: Document,
    placeholder: str,
    image_urls: List[str],
    language: str = "zh",
    *,
    replace_in_paragraph_fn: Optional[ReplaceInParagraphFn] = None,
    replace_in_cell_fn: Optional[ReplaceInCellFn] = None,
    insert_block_paragraph_fn: Optional[InsertBlockParagraphFn] = None,
    insert_cell_paragraph_fn: Optional[InsertCellParagraphFn] = None,
) -> bool:
    """定位占位符并用图片替换，返回是否找到并处理了占位符。

    Step 1: 在正文中查找占位符
    Step 2: 正文未命中则在表格单元格中查找
    Step 3: 表格未命中则在页眉/页脚中查找
    Step 4: 找到后委托 _apply_images_to_* 执行「下载 → 删占位符/兜底」
    """
    # Step 1: 正文段落
    if _apply_images_in_paragraphs(
        doc.paragraphs,
        doc,
        placeholder,
        image_urls,
        language,
        replace_in_paragraph_fn=replace_in_paragraph_fn,
        insert_block_paragraph_fn=insert_block_paragraph_fn,
    ):
        return True

    # Step 2: 表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder not in (cell.text or ""):
                    continue
                _apply_images_to_cell(
                    cell,
                    placeholder,
                    image_urls,
                    language,
                    replace_in_cell_fn=replace_in_cell_fn,
                    insert_cell_paragraph_fn=insert_cell_paragraph_fn,
                )
                return True

    # Step 3: 页眉 / 页脚
    for section in doc.sections:
        for paragraphs in (
            section.header.paragraphs if section.header else [],
            section.footer.paragraphs if section.footer else [],
        ):
            if _apply_images_in_paragraphs(
                paragraphs,
                doc,
                placeholder,
                image_urls,
                language,
                replace_in_paragraph_fn=replace_in_paragraph_fn,
                insert_block_paragraph_fn=insert_block_paragraph_fn,
            ):
                return True

    return False